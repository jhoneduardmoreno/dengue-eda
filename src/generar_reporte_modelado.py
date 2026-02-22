"""
Generador de Reporte Word: Modelado Predictivo de Exceso Epidemico - Dengue.

Replica la logica del notebook 06_modelado.ipynb, entrena 3 modelos
(Logistic Regression, Random Forest, XGBoost), genera graficos PNG
y ensambla un documento Word profesional con los resultados.

Ejecucion:
    conda run -n dengue-eda python src/generar_reporte_modelado.py

Genera:
    reporte_modelado_dengue.docx (en la raiz del proyecto)
    results_graphs/modelado/*.png (graficos individuales)
"""

import warnings
warnings.filterwarnings("ignore")

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path

from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import (
    classification_report, confusion_matrix, roc_auc_score, roc_curve,
    f1_score, precision_score, recall_score, accuracy_score,
)
from xgboost import XGBClassifier
import joblib

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================================
# Rutas y constantes
# ============================================================================
PROJECT_ROOT = Path(__file__).resolve().parent.parent
DATA_PATH = PROJECT_ROOT / "data" / "processed" / "panel_municipal_mensual.parquet"
GRAPHS_DIR = PROJECT_ROOT / "results_graphs" / "modelado"
OUTPUT_DOCX = PROJECT_ROOT / "reporte_modelado_dengue.docx"
MODELS_DIR = PROJECT_ROOT / "models"

# Colores para graficos
COLOR_LR = "#2196F3"
COLOR_RF = "#4CAF50"
COLOR_XGB = "#FF9800"

# Textos del reporte
TEXTOS = {
    "titulo": "Modelado Predictivo de\nExceso Epidemico de Dengue",
    "universidad": "Universidad Autonoma de Occidente",
    "maestria": "Maestria en Inteligencia Artificial",
    "materia": "Desarrollo de Soluciones con IA",
    "autores": [
        "Jhon Edwar Salazar",
        "Santiago Castano Orozco",
        "David Alejandro Burbano Getial",
    ],
    "fecha": "Febrero 2025",
}

# Features
FEATS_CLIMA = ["temperatura_c", "precipitacion_mm", "ndvi", "dewpoint_c"]
FEATS_CLIMA_LAGS = [f"{v}_lag{l}" for v in FEATS_CLIMA for l in [1, 2, 3]]
FEATS_CLIMA_MM = [f"{v}_mm3" for v in FEATS_CLIMA]
FEATS_EPI_LAGS = [
    "casos_total_lag1", "casos_total_lag2", "casos_total_lag3",
    "tasa_incidencia_lag1", "tasa_incidencia_lag2", "tasa_incidencia_lag3",
]
FEATS_EPI_MM = ["casos_total_mm3", "tasa_incidencia_mm3"]
FEATS_DEMO = ["prop_grave", "prop_hospitalizado", "prop_femenino", "poblacion"]

FEATURES = FEATS_CLIMA + FEATS_CLIMA_LAGS + FEATS_CLIMA_MM + FEATS_EPI_LAGS + FEATS_EPI_MM + FEATS_DEMO
TARGET = "exceso"

ANOS_TRAIN = [2010, 2016, 2019, 2022]
ANOS_TEST = [2024]


# ============================================================================
# Pipeline de datos y modelado
# ============================================================================
def cargar_y_preparar_datos():
    """Carga el parquet, selecciona features, dropea NaN y split temporal."""
    print("  Cargando datos...")
    df = pd.read_parquet(DATA_PATH)
    print(f"    Dimensiones originales: {df.shape[0]:,} x {df.shape[1]}")

    df_model = df[FEATURES + [TARGET, "ano"]].copy()
    n_antes = len(df_model)
    df_model = df_model.dropna(subset=FEATURES)
    print(f"    Despues de dropear NaN: {len(df_model):,} filas ({len(df_model)/n_antes*100:.1f}%)")

    # Split temporal
    train_mask = df_model["ano"].isin(ANOS_TRAIN)
    test_mask = df_model["ano"].isin(ANOS_TEST)

    X_train = df_model.loc[train_mask, FEATURES]
    y_train = df_model.loc[train_mask, TARGET]
    X_test = df_model.loc[test_mask, FEATURES]
    y_test = df_model.loc[test_mask, TARGET]

    print(f"    Train: {len(X_train):,} filas | Test: {len(X_test):,} filas")

    return df, df_model, X_train, y_train, X_test, y_test


def entrenar_modelos(X_train, y_train, X_test):
    """Entrena los 3 modelos y retorna predicciones."""
    # Logistic Regression (necesita escalado)
    print("  Entrenando Logistic Regression...")
    scaler = StandardScaler()
    X_train_scaled = scaler.fit_transform(X_train)
    X_test_scaled = scaler.transform(X_test)

    lr = LogisticRegression(
        class_weight="balanced", max_iter=1000, random_state=42, solver="lbfgs"
    )
    lr.fit(X_train_scaled, y_train)
    y_pred_lr = lr.predict(X_test_scaled)
    y_prob_lr = lr.predict_proba(X_test_scaled)[:, 1]

    # Random Forest
    print("  Entrenando Random Forest...")
    rf = RandomForestClassifier(
        n_estimators=200, class_weight="balanced", random_state=42, n_jobs=-1
    )
    rf.fit(X_train, y_train)
    y_pred_rf = rf.predict(X_test)
    y_prob_rf = rf.predict_proba(X_test)[:, 1]

    # XGBoost
    print("  Entrenando XGBoost...")
    n_neg = (y_train == 0).sum()
    n_pos = (y_train == 1).sum()
    spw = n_neg / n_pos

    xgb = XGBClassifier(
        n_estimators=200, max_depth=6, learning_rate=0.1,
        scale_pos_weight=spw, random_state=42, eval_metric="logloss", n_jobs=-1,
    )
    xgb.fit(X_train, y_train)
    y_pred_xgb = xgb.predict(X_test)
    y_prob_xgb = xgb.predict_proba(X_test)[:, 1]

    modelos = {
        "Logistic Regression": {"pred": y_pred_lr, "prob": y_prob_lr, "model": lr, "scaler": scaler},
        "Random Forest": {"pred": y_pred_rf, "prob": y_prob_rf, "model": rf},
        "XGBoost": {"pred": y_pred_xgb, "prob": y_prob_xgb, "model": xgb},
    }
    return modelos


def exportar_modelo_lr(modelos):
    """Exporta el modelo Logistic Regression (con scaler y features) a joblib."""
    MODELS_DIR.mkdir(parents=True, exist_ok=True)
    lr_data = modelos["Logistic Regression"]
    artefacto = {
        "model": lr_data["model"],
        "scaler": lr_data["scaler"],
        "features": FEATURES,
    }
    ruta = MODELS_DIR / "logistic_regression.joblib"
    joblib.dump(artefacto, ruta)
    tamano_kb = ruta.stat().st_size / 1024
    print(f"  Modelo exportado: {ruta}  ({tamano_kb:.1f} KB)")
    return ruta


def calcular_metricas(y_test, modelos):
    """Calcula metricas para cada modelo."""
    resultados = []
    for nombre, m in modelos.items():
        resultados.append({
            "Modelo": nombre,
            "Accuracy": accuracy_score(y_test, m["pred"]),
            "Precision": precision_score(y_test, m["pred"]),
            "Recall": recall_score(y_test, m["pred"]),
            "F1-Score": f1_score(y_test, m["pred"]),
            "ROC-AUC": roc_auc_score(y_test, m["prob"]),
        })
    return pd.DataFrame(resultados).set_index("Modelo")


# ============================================================================
# Generacion de graficos
# ============================================================================
def configurar_estilo():
    sns.set_theme(style="whitegrid", palette="muted", font_scale=1.1)
    plt.rcParams.update({
        "figure.dpi": 150,
        "axes.titlesize": 14,
        "axes.labelsize": 12,
        "savefig.bbox": "tight",
        "savefig.dpi": 150,
    })


def grafico_balance_target(df, df_model):
    """Balance de variable objetivo: global + por ano."""
    fig, axes = plt.subplots(1, 2, figsize=(14, 5))

    counts = df_model[TARGET].value_counts()
    bars = axes[0].bar(
        ["Sin exceso (0)", "Con exceso (1)"], counts.values,
        color=[COLOR_LR, "#F44336"], edgecolor="white",
    )
    for bar, val in zip(bars, counts.values):
        axes[0].text(
            bar.get_x() + bar.get_width() / 2, bar.get_height() + 200,
            f"{val:,}\n({val / len(df_model) * 100:.1f}%)",
            ha="center", va="bottom", fontweight="bold",
        )
    axes[0].set_title("Balance de la variable objetivo", fontweight="bold")
    axes[0].set_ylabel("Numero de observaciones")

    balance_ano = df_model.groupby("ano")[TARGET].mean().mul(100)
    bars2 = axes[1].bar(
        balance_ano.index.astype(str), balance_ano.values,
        color="#F44336", edgecolor="white", alpha=0.8,
    )
    for bar, val in zip(bars2, balance_ano.values):
        axes[1].text(
            bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.2,
            f"{val:.1f}%", ha="center", va="bottom", fontweight="bold",
        )
    axes[1].set_title("Porcentaje de exceso epidemico por ano", fontweight="bold")
    axes[1].set_ylabel("% municipios-mes con exceso")
    axes[1].set_xlabel("Ano")

    plt.tight_layout()
    ruta = GRAPHS_DIR / "balance_target.png"
    fig.savefig(ruta)
    plt.close(fig)
    return ruta


def grafico_confusion_matrix(y_test, y_pred, nombre, color_map, filename):
    """Genera confusion matrix para un modelo."""
    fig, ax = plt.subplots(figsize=(7, 5))
    cm = confusion_matrix(y_test, y_pred)
    sns.heatmap(
        cm, annot=True, fmt="d", cmap=color_map, ax=ax,
        xticklabels=["Sin exceso", "Con exceso"],
        yticklabels=["Sin exceso", "Con exceso"],
    )
    ax.set_title(f"Matriz de Confusion - {nombre}", fontweight="bold")
    ax.set_ylabel("Real")
    ax.set_xlabel("Predicho")
    plt.tight_layout()
    ruta = GRAPHS_DIR / filename
    fig.savefig(ruta)
    plt.close(fig)
    return ruta


def grafico_roc_individual(y_test, y_prob, nombre, color, filename):
    """Genera curva ROC para un modelo."""
    fig, ax = plt.subplots(figsize=(7, 5))
    fpr, tpr, _ = roc_curve(y_test, y_prob)
    auc = roc_auc_score(y_test, y_prob)
    ax.plot(fpr, tpr, color=color, lw=2, label=f"{nombre} (AUC = {auc:.4f})")
    ax.plot([0, 1], [0, 1], "k--", lw=1, alpha=0.5)
    ax.set_xlabel("False Positive Rate")
    ax.set_ylabel("True Positive Rate")
    ax.set_title(f"Curva ROC - {nombre}", fontweight="bold")
    ax.legend(loc="lower right")
    plt.tight_layout()
    ruta = GRAPHS_DIR / filename
    fig.savefig(ruta)
    plt.close(fig)
    return ruta


def grafico_feature_importance(model, nombre, color, filename, top_n=15):
    """Genera feature importance top N."""
    fig, ax = plt.subplots(figsize=(8, 6))
    importances = pd.Series(model.feature_importances_, index=FEATURES)
    top = importances.nlargest(top_n)
    top.sort_values().plot(kind="barh", ax=ax, color=color, edgecolor="white")
    ax.set_title(f"Top {top_n} Features - {nombre}", fontweight="bold")
    ax.set_xlabel("Importancia")
    plt.tight_layout()
    ruta = GRAPHS_DIR / filename
    fig.savefig(ruta)
    plt.close(fig)
    return ruta


def grafico_comparacion_metricas(df_resultados):
    """Barras comparativas de metricas."""
    fig, ax = plt.subplots(figsize=(10, 6))
    metricas_plot = df_resultados[["F1-Score", "ROC-AUC", "Recall", "Precision"]]
    metricas_plot.plot(kind="bar", ax=ax, edgecolor="white", width=0.7)
    ax.set_title("Comparacion de Metricas por Modelo", fontweight="bold")
    ax.set_ylabel("Score")
    ax.set_ylim(0, 1)
    ax.legend(loc="upper left", fontsize=9)
    ax.set_xticklabels(ax.get_xticklabels(), rotation=0)
    plt.tight_layout()
    ruta = GRAPHS_DIR / "comparacion_metricas.png"
    fig.savefig(ruta)
    plt.close(fig)
    return ruta


def grafico_roc_comparativo(y_test, modelos):
    """Curvas ROC superpuestas."""
    fig, ax = plt.subplots(figsize=(8, 6))
    colores = {
        "Logistic Regression": COLOR_LR,
        "Random Forest": COLOR_RF,
        "XGBoost": COLOR_XGB,
    }
    for nombre, m in modelos.items():
        fpr, tpr, _ = roc_curve(y_test, m["prob"])
        auc = roc_auc_score(y_test, m["prob"])
        ax.plot(fpr, tpr, color=colores[nombre], lw=2,
                label=f"{nombre} (AUC={auc:.4f})")
    ax.plot([0, 1], [0, 1], "k--", lw=1, alpha=0.5)
    ax.set_xlabel("False Positive Rate")
    ax.set_ylabel("True Positive Rate")
    ax.set_title("Curvas ROC Comparativas", fontweight="bold")
    ax.legend(loc="lower right")
    plt.tight_layout()
    ruta = GRAPHS_DIR / "roc_comparativo.png"
    fig.savefig(ruta)
    plt.close(fig)
    return ruta


def generar_todos_los_graficos(df, df_model, y_test, modelos, df_resultados):
    """Genera todos los graficos y retorna dict con rutas."""
    GRAPHS_DIR.mkdir(parents=True, exist_ok=True)

    graficos = {}
    print("  Generando graficos...")

    graficos["balance_target"] = grafico_balance_target(df, df_model)
    print("    [1/11] balance_target.png")

    # Logistic Regression
    graficos["cm_logistic"] = grafico_confusion_matrix(
        y_test, modelos["Logistic Regression"]["pred"],
        "Logistic Regression", "Blues", "cm_logistic.png",
    )
    print("    [2/11] cm_logistic.png")

    graficos["roc_logistic"] = grafico_roc_individual(
        y_test, modelos["Logistic Regression"]["prob"],
        "Logistic Regression", COLOR_LR, "roc_logistic.png",
    )
    print("    [3/11] roc_logistic.png")

    # Random Forest
    graficos["cm_random_forest"] = grafico_confusion_matrix(
        y_test, modelos["Random Forest"]["pred"],
        "Random Forest", "Greens", "cm_random_forest.png",
    )
    print("    [4/11] cm_random_forest.png")

    graficos["roc_random_forest"] = grafico_roc_individual(
        y_test, modelos["Random Forest"]["prob"],
        "Random Forest", COLOR_RF, "roc_random_forest.png",
    )
    print("    [5/11] roc_random_forest.png")

    graficos["fi_random_forest"] = grafico_feature_importance(
        modelos["Random Forest"]["model"], "Random Forest", COLOR_RF,
        "fi_random_forest.png",
    )
    print("    [6/11] fi_random_forest.png")

    # XGBoost
    graficos["cm_xgboost"] = grafico_confusion_matrix(
        y_test, modelos["XGBoost"]["pred"],
        "XGBoost", "Oranges", "cm_xgboost.png",
    )
    print("    [7/11] cm_xgboost.png")

    graficos["roc_xgboost"] = grafico_roc_individual(
        y_test, modelos["XGBoost"]["prob"],
        "XGBoost", COLOR_XGB, "roc_xgboost.png",
    )
    print("    [8/11] roc_xgboost.png")

    graficos["fi_xgboost"] = grafico_feature_importance(
        modelos["XGBoost"]["model"], "XGBoost", COLOR_XGB,
        "fi_xgboost.png",
    )
    print("    [9/11] fi_xgboost.png")

    # Comparativos
    graficos["comparacion_metricas"] = grafico_comparacion_metricas(df_resultados)
    print("    [10/11] comparacion_metricas.png")

    graficos["roc_comparativo"] = grafico_roc_comparativo(y_test, modelos)
    print("    [11/11] roc_comparativo.png")

    return graficos


# ============================================================================
# Generacion del documento Word
# ============================================================================
IMG_WIDTH = Cm(16)

# Colores Word
AZUL_OSCURO = RGBColor(0, 71, 133)
AZUL_MEDIO = RGBColor(41, 128, 185)
BLANCO = RGBColor(255, 255, 255)
GRIS = RGBColor(120, 120, 120)
NEGRO = RGBColor(33, 33, 33)


def set_cell_shading(cell, hex_color):
    """Aplica color de fondo a una celda."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="000000", sz="4"):
    """Aplica bordes finos a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f"</w:tcBorders>"
    )
    tcPr.append(borders)


def agregar_tabla_formateada(doc, headers, rows):
    """Agrega una tabla con header azul y bordes."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = BLANCO
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "004789")
        set_cell_borders(cell, "003366")

    # Filas
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = NEGRO
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bg = "F5F8FC" if r_idx % 2 == 0 else "FFFFFF"
            set_cell_shading(cell, bg)
            set_cell_borders(cell, "CCCCCC")

    doc.add_paragraph("")
    return table


def agregar_portada(doc):
    """Agrega portada al documento."""
    # Espaciado superior
    for _ in range(4):
        doc.add_paragraph("")

    # Titulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in TEXTOS["titulo"].split("\n"):
        run = p.add_run(line + "\n")
        run.bold = True
        run.font.size = Pt(26)
        run.font.color.rgb = AZUL_OSCURO

    # Linea decorativa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.color.rgb = AZUL_MEDIO

    doc.add_paragraph("")

    # Universidad, maestria, materia
    for texto, size, bold in [
        (TEXTOS["universidad"], 14, False),
        (TEXTOS["maestria"], 12, False),
        (TEXTOS["materia"], 12, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(texto)
        run.font.size = Pt(size)
        run.font.color.rgb = NEGRO
        if bold:
            run.bold = True

    doc.add_paragraph("")

    # Autores
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Autores:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NEGRO

    for autor in TEXTOS["autores"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(autor)
        run.font.size = Pt(11)
        run.font.color.rgb = NEGRO

    doc.add_paragraph("")

    # Fecha
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TEXTOS["fecha"])
    run.font.size = Pt(11)
    run.font.color.rgb = GRIS

    doc.add_page_break()


def agregar_imagen(doc, ruta, width=IMG_WIDTH):
    """Agrega una imagen centrada."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(ruta), width=width)
    doc.add_paragraph("")


def agregar_classification_report_tabla(doc, y_test, y_pred, nombre):
    """Agrega el classification report como tabla Word."""
    report = classification_report(y_test, y_pred, target_names=["Sin exceso", "Con exceso"], output_dict=True)

    headers = ["", "Precision", "Recall", "F1-Score", "Support"]
    rows = []
    for cls_name in ["Sin exceso", "Con exceso"]:
        r = report[cls_name]
        rows.append([
            cls_name,
            f"{r['precision']:.4f}",
            f"{r['recall']:.4f}",
            f"{r['f1-score']:.4f}",
            f"{int(r['support']):,}",
        ])

    # Agregar accuracy, macro avg, weighted avg
    rows.append([
        "Accuracy",
        "", "",
        f"{report['accuracy']:.4f}",
        f"{int(report['weighted avg']['support']):,}",
    ])
    for avg_name in ["macro avg", "weighted avg"]:
        r = report[avg_name]
        rows.append([
            avg_name.title(),
            f"{r['precision']:.4f}",
            f"{r['recall']:.4f}",
            f"{r['f1-score']:.4f}",
            f"{int(r['support']):,}",
        ])

    agregar_tabla_formateada(doc, headers, rows)


def generar_documento_word(df, df_model, X_train, y_train, X_test, y_test,
                           modelos, df_resultados, graficos):
    """Genera el documento Word completo."""
    doc = Document()

    # Configurar estilo base
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = NEGRO

    # ── Portada ──
    print("  [1/8] Portada...")
    agregar_portada(doc)

    # ── Seccion 1: Introduccion ──
    print("  [2/8] Introduccion...")
    doc.add_heading("1. Introduccion", level=1)
    doc.add_paragraph(
        "Este reporte presenta los resultados del modelado predictivo para "
        "la clasificacion de exceso epidemico de dengue a nivel municipal-mensual "
        "en Colombia. Se entrenaron tres modelos de clasificacion supervisada "
        "(Logistic Regression, Random Forest y XGBoost) utilizando un panel de "
        "datos que integra variables climaticas, epidemiologicas y demograficas."
    )
    doc.add_paragraph(
        "El objetivo es predecir si un municipio experimentara un exceso epidemico "
        "de dengue (variable binaria) en un mes dado, definido como una incidencia "
        "superior al umbral historico. Los modelos fueron entrenados con datos de "
        "los anos 2010, 2016, 2019 y 2022, y evaluados en datos del ano 2024."
    )

    # ── Seccion 2: Datos y features ──
    print("  [3/8] Datos y features...")
    doc.add_heading("2. Datos y Features", level=1)
    doc.add_paragraph(
        f"Se utilizo el panel municipal-mensual (panel_municipal_mensual.parquet) "
        f"con {len(df):,} observaciones originales. Tras eliminar filas con "
        f"valores nulos en los features (generados por lags temporales), "
        f"se trabajaron {len(df_model):,} observaciones."
    )

    doc.add_heading("2.1 Features seleccionados", level=2)
    doc.add_paragraph(
        f"Se seleccionaron {len(FEATURES)} features agrupados en las siguientes categorias:"
    )

    feature_groups = [
        ("Climaticos actuales", FEATS_CLIMA),
        ("Climaticos lags (1-3 meses)", FEATS_CLIMA_LAGS),
        ("Climaticos media movil 3 meses", FEATS_CLIMA_MM),
        ("Epidemiologicos lags (1-3 meses)", FEATS_EPI_LAGS),
        ("Epidemiologicos media movil 3 meses", FEATS_EPI_MM),
        ("Demograficos y poblacion", FEATS_DEMO),
    ]
    headers_feat = ["Grupo", "Cantidad", "Features"]
    rows_feat = []
    for grupo, feats in feature_groups:
        rows_feat.append([grupo, str(len(feats)), ", ".join(feats)])
    agregar_tabla_formateada(doc, headers_feat, rows_feat)

    doc.add_heading("2.2 Balance de la variable objetivo", level=2)
    counts = df_model[TARGET].value_counts()
    ratio = counts[0] / counts[1]
    doc.add_paragraph(
        f"La variable objetivo 'exceso' esta altamente desbalanceada: "
        f"{counts[0]:,} observaciones sin exceso ({counts[0]/len(df_model)*100:.1f}%) vs "
        f"{counts[1]:,} con exceso ({counts[1]/len(df_model)*100:.1f}%). "
        f"Ratio de desbalance: {ratio:.1f}:1. "
        f"Se utilizaron tecnicas de manejo de desbalance: class_weight='balanced' "
        f"(Logistic Regression y Random Forest) y scale_pos_weight (XGBoost)."
    )
    agregar_imagen(doc, graficos["balance_target"])

    # ── Seccion 3: Split temporal ──
    print("  [4/8] Split temporal...")
    doc.add_heading("3. Split Temporal Train/Test", level=1)
    doc.add_paragraph(
        f"Se utilizo un split temporal para respetar la naturaleza secuencial de los datos:"
    )

    headers_split = ["Conjunto", "Anos", "Filas", "% Exceso"]
    balance_train = (y_train == 1).mean() * 100
    balance_test = (y_test == 1).mean() * 100
    rows_split = [
        ["Train", ", ".join(str(a) for a in ANOS_TRAIN), f"{len(X_train):,}", f"{balance_train:.1f}%"],
        ["Test", ", ".join(str(a) for a in ANOS_TEST), f"{len(X_test):,}", f"{balance_test:.1f}%"],
    ]
    agregar_tabla_formateada(doc, headers_split, rows_split)

    # ── Seccion 4: Logistic Regression ──
    print("  [5/8] Logistic Regression...")
    doc.add_heading("4. Logistic Regression", level=1)
    doc.add_paragraph(
        "Modelo lineal con regularizacion L2 (solver='lbfgs'), class_weight='balanced' "
        "para compensar el desbalance de clases. Los features fueron estandarizados "
        "con StandardScaler. Maximo de 1000 iteraciones."
    )
    doc.add_heading("Classification Report", level=2)
    agregar_classification_report_tabla(doc, y_test, modelos["Logistic Regression"]["pred"], "Logistic Regression")

    doc.add_heading("Matriz de Confusion", level=2)
    agregar_imagen(doc, graficos["cm_logistic"])

    doc.add_heading("Curva ROC", level=2)
    agregar_imagen(doc, graficos["roc_logistic"])

    # ── Seccion 5: Random Forest ──
    print("  [6/8] Random Forest...")
    doc.add_heading("5. Random Forest", level=1)
    doc.add_paragraph(
        "Ensemble de 200 arboles de decision con class_weight='balanced'. "
        "No requiere escalado de features. Utiliza paralelizacion (n_jobs=-1)."
    )
    doc.add_heading("Classification Report", level=2)
    agregar_classification_report_tabla(doc, y_test, modelos["Random Forest"]["pred"], "Random Forest")

    doc.add_heading("Matriz de Confusion", level=2)
    agregar_imagen(doc, graficos["cm_random_forest"])

    doc.add_heading("Curva ROC", level=2)
    agregar_imagen(doc, graficos["roc_random_forest"])

    doc.add_heading("Feature Importance (Top 15)", level=2)
    agregar_imagen(doc, graficos["fi_random_forest"])

    # ── Seccion 6: XGBoost ──
    print("  [7/8] XGBoost...")
    doc.add_heading("6. XGBoost", level=1)

    n_neg = (y_train == 0).sum()
    n_pos = (y_train == 1).sum()
    spw = n_neg / n_pos
    doc.add_paragraph(
        f"Gradient Boosting con 200 estimadores, max_depth=6, learning_rate=0.1. "
        f"Manejo de desbalance mediante scale_pos_weight={spw:.2f} "
        f"(ratio negativo/positivo). Metrica de evaluacion: logloss."
    )
    doc.add_heading("Classification Report", level=2)
    agregar_classification_report_tabla(doc, y_test, modelos["XGBoost"]["pred"], "XGBoost")

    doc.add_heading("Matriz de Confusion", level=2)
    agregar_imagen(doc, graficos["cm_xgboost"])

    doc.add_heading("Curva ROC", level=2)
    agregar_imagen(doc, graficos["roc_xgboost"])

    doc.add_heading("Feature Importance (Top 15)", level=2)
    agregar_imagen(doc, graficos["fi_xgboost"])

    # ── Seccion 7: Comparacion de modelos ──
    print("  [8/8] Comparacion y conclusiones...")
    doc.add_heading("7. Comparacion de Modelos", level=1)
    doc.add_paragraph(
        "Tabla resumen de metricas de los tres modelos evaluados sobre el "
        "conjunto de test (ano 2024):"
    )

    headers_comp = ["Modelo", "Accuracy", "Precision", "Recall", "F1-Score", "ROC-AUC"]
    rows_comp = []
    for nombre in df_resultados.index:
        row = df_resultados.loc[nombre]
        rows_comp.append([
            nombre,
            f"{row['Accuracy']:.4f}",
            f"{row['Precision']:.4f}",
            f"{row['Recall']:.4f}",
            f"{row['F1-Score']:.4f}",
            f"{row['ROC-AUC']:.4f}",
        ])
    agregar_tabla_formateada(doc, headers_comp, rows_comp)

    doc.add_heading("Comparacion grafica de metricas", level=2)
    agregar_imagen(doc, graficos["comparacion_metricas"])

    doc.add_heading("Curvas ROC comparativas", level=2)
    agregar_imagen(doc, graficos["roc_comparativo"])

    # ── Seccion 8: Conclusiones ──
    doc.add_heading("8. Conclusiones", level=1)

    mejor_f1 = df_resultados["F1-Score"].idxmax()
    mejor_auc = df_resultados["ROC-AUC"].idxmax()

    doc.add_paragraph(
        f"1. Mejor modelo por F1-Score: {mejor_f1} "
        f"({df_resultados.loc[mejor_f1, 'F1-Score']:.4f})"
    )
    doc.add_paragraph(
        f"2. Mejor modelo por ROC-AUC: {mejor_auc} "
        f"({df_resultados.loc[mejor_auc, 'ROC-AUC']:.4f})"
    )
    doc.add_paragraph(
        "3. El dataset esta altamente desbalanceado (~5% positivo), por lo que las "
        "tecnicas de class_weight/scale_pos_weight son fundamentales para obtener "
        "recall aceptable en la clase minoritaria."
    )
    doc.add_paragraph(
        "4. Las metricas principales para evaluar el desempeno son F1-Score y "
        "ROC-AUC, ya que accuracy no es informativa con clases tan desbalanceadas."
    )
    doc.add_paragraph(
        "5. Los features mas importantes incluyen variables epidemiologicas con rezago "
        "temporal (casos y tasas de incidencia en meses previos), lo que confirma la "
        "naturaleza autoregresiva de los brotes de dengue."
    )

    # Guardar
    doc.save(str(OUTPUT_DOCX))
    return OUTPUT_DOCX


# ============================================================================
# Funcion principal
# ============================================================================
def main():
    print("=" * 60)
    print("  Generador de Reporte de Modelado - Dengue Colombia")
    print("=" * 60)

    # 1. Datos
    print("\n[1/4] Cargando y preparando datos...")
    df, df_model, X_train, y_train, X_test, y_test = cargar_y_preparar_datos()

    # 2. Modelos
    print("\n[2/4] Entrenando modelos...")
    modelos = entrenar_modelos(X_train, y_train, X_test)
    exportar_modelo_lr(modelos)

    # 3. Metricas
    print("\n[3/4] Calculando metricas y generando graficos...")
    df_resultados = calcular_metricas(y_test, modelos)
    print("\n  Resumen de metricas:")
    print(df_resultados.round(4).to_string())

    configurar_estilo()
    graficos = generar_todos_los_graficos(df, df_model, y_test, modelos, df_resultados)

    # 4. Documento Word
    print("\n[4/4] Generando documento Word...")
    ruta_docx = generar_documento_word(
        df, df_model, X_train, y_train, X_test, y_test,
        modelos, df_resultados, graficos,
    )

    print(f"\n{'=' * 60}")
    print(f"  Reporte generado: {ruta_docx}")
    print(f"  Tamano: {ruta_docx.stat().st_size / 1024:.1f} KB")
    print(f"  Graficos en: {GRAPHS_DIR}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
