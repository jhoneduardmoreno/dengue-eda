"""
Genera un reporte Word consolidado del proyecto dengue cubriendo EDA + Modelado
+ decisiones + resultados. Usa los gráficos ya generados en results_graphs/foco/
y lee panel + predicciones del modelado per-municipio.

Salida: docs/reportes/reporte_proyecto_dengue_<YYYY-MM-DD>.docx

Reusa los helpers de formato Word de src/generar_reporte_modelado.py
(Entrega 2): tablas con header azul, imágenes centradas, colores corporativos.

Uso:
    conda run -n dengue-eda python scripts/generar_reporte_consolidado.py
"""
from __future__ import annotations

import os
import sys
import warnings
from datetime import datetime
from pathlib import Path

warnings.filterwarnings("ignore")

import joblib
import numpy as np
import pandas as pd
from sklearn.metrics import (
    accuracy_score,
    confusion_matrix,
    f1_score,
    precision_score,
    recall_score,
)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))
from utils import MUNICIPIOS_FOCO, PROJECT_ROOT  # noqa: E402
from generar_reporte_modelado import (  # noqa: E402
    AZUL_MEDIO,
    AZUL_OSCURO,
    BLANCO,
    GRIS,
    IMG_WIDTH,
    NEGRO,
    agregar_imagen,
    agregar_tabla_formateada,
)

# ============================================================================
# Rutas y constantes
# ============================================================================
PANEL_PATH = PROJECT_ROOT / "data" / "processed" / "panel_municipal_mensual.parquet"
PRED_PATH = PROJECT_ROOT / "data" / "processed" / "predicciones_test.csv"
SENS_PATH = PROJECT_ROOT / "data" / "processed" / "sensibilidad_target.csv"
MODELS_DIR = PROJECT_ROOT / "models"
GRAPHS_DIR = PROJECT_ROOT / "results_graphs" / "foco"
OUT_DIR = PROJECT_ROOT / "docs" / "reportes"

TEXTOS = {
    "titulo": ("Sistema de alerta temprana de excesos epidémicos de dengue\n"
               "Modelos per-municipio para Colombia (2007–2024)"),
    "universidad": "Universidad de los Andes",
    "maestria": "Maestría en Inteligencia Artificial",
    "materia": "Desarrollo de Soluciones",
    "autores": [
        "Jhon Edward Moreno Díaz",
        "Hernán",
        "Ruby",
        "Danilo",
    ],
    "fecha": "Mayo 2026",
}

DPTO_POR_MPIO = {
    "23855": ("Valencia", "Córdoba", "Caribe"),
    "47288": ("Fundación", "Magdalena", "Sierra Nevada"),
    "95025": ("El Retorno", "Guaviare", "Amazonía"),
}

ANO_TEST_DESDE = 2020


# ============================================================================
# Helpers locales
# ============================================================================
def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = AZUL_OSCURO
    return h


def parrafo(doc, text, *, justify=True, bold=False, size=11):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = NEGRO
    if bold:
        run.bold = True
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = NEGRO


def agregar_portada(doc, textos):
    """Portada del proyecto con los textos del proyecto actual."""
    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in textos["titulo"].split("\n"):
        run = p.add_run(line + "\n")
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = AZUL_OSCURO

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.color.rgb = AZUL_MEDIO

    doc.add_paragraph("")
    for texto, size in [
        (textos["universidad"], 14),
        (textos["maestria"], 12),
        (textos["materia"], 12),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(texto)
        run.font.size = Pt(size)
        run.font.color.rgb = NEGRO

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Autores")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NEGRO

    for autor in textos["autores"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(autor)
        run.font.size = Pt(11)
        run.font.color.rgb = NEGRO

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(textos["fecha"])
    run.font.size = Pt(11)
    run.font.color.rgb = GRIS

    doc.add_page_break()


def metricas_modelo(pred, cod_mpio: int, col_pred: str):
    """Retorna dict con P/R/F1/Accuracy/TP/FP/TN/FN."""
    sub = pred[pred["cod_mpio"] == cod_mpio]
    y, yhat = sub["exceso"].astype(int).values, sub[col_pred].astype(int).values
    cm = confusion_matrix(y, yhat, labels=[0, 1])
    tn, fp, fn, tp = int(cm[0, 0]), int(cm[0, 1]), int(cm[1, 0]), int(cm[1, 1])
    return {
        "precision": precision_score(y, yhat, zero_division=0),
        "recall":    recall_score(y, yhat, zero_division=0),
        "f1":        f1_score(y, yhat, zero_division=0),
        "accuracy":  accuracy_score(y, yhat),
        "TP": tp, "FP": fp, "TN": tn, "FN": fn,
    }


# ============================================================================
# Secciones del reporte
# ============================================================================
def seccion_resumen(doc, pred):
    heading(doc, "1. Resumen ejecutivo")

    parrafo(doc,
        "Este proyecto desarrolla un sistema de alerta temprana de excesos "
        "epidemiológicos de dengue a nivel municipal para Colombia. Evoluciona "
        "el modelo nacional único de la Entrega 1 (Precision = 0.43, Recall = "
        "0.87) hacia modelos independientes por municipio, atendiendo la "
        "recomendación del director del proyecto en la reunión del 24 de abril "
        "de 2026: \"asuman que cada unidad espacial va a ser distinta y el "
        "modelo va a ajustarse de manera diferente\".")

    parrafo(doc,
        "Se eligieron tres municipios foco bajo el criterio de mayor incidencia "
        "por habitante, priorizando diversidad climática y regional: Valencia "
        "(Córdoba, Caribe), Fundación (Magdalena, Sierra Nevada) y El Retorno "
        "(Guaviare, Amazonía). Se entrenaron modelos de Regresión Logística "
        "regularizada y XGBoost con partición temporal estricta (train "
        "2007–2019, test 2020–2024), target binario de exceso definido por "
        "percentil 75 histórico mensual con piso de 2 casos, y un baseline "
        "trivial \"casos del mes anterior > 2\" como piso comparativo.")

    # headline metrics
    m_v = metricas_modelo(pred, 23855, "pred_xgboost")
    m_f = metricas_modelo(pred, 47288, "pred_xgboost")
    m_r = metricas_modelo(pred, 95025, "pred_xgboost")
    parrafo(doc,
        f"Resultado principal: en Valencia, XGBoost obtiene Precision = "
        f"{m_v['precision']:.2f} y Recall = {m_v['recall']:.2f}; en Fundación, "
        f"Precision = {m_f['precision']:.2f} y Recall = {m_f['recall']:.2f}. "
        f"Ambos casi duplican la Precision del modelo nacional original "
        f"manteniendo Recall alto, cumpliendo el objetivo cuantitativo del "
        f"proyecto. En El Retorno, ningún modelo entrenado supera al baseline "
        f"trivial (XGBoost: Precision = {m_r['precision']:.2f}, Recall = "
        f"{m_r['recall']:.2f}); se documenta como limitación intrínseca de la "
        f"serie y como hallazgo metodológico — la aproximación regional tiene "
        f"un piso de aplicabilidad relacionado con la consistencia de los datos.")

    doc.add_page_break()


def seccion_contexto(doc):
    heading(doc, "2. Contexto y objetivo")

    parrafo(doc,
        "El dengue es la enfermedad transmitida por vectores con mayor carga "
        "epidemiológica en Colombia. Los brotes presentan dinámica multifactorial "
        "que combina factores climáticos (temperatura, precipitación, humedad), "
        "demográficos y sociales, y muestran heterogeneidad regional fuerte "
        "que dificulta construir un modelo nacional único confiable.")

    parrafo(doc,
        "El proyecto parte de una versión previa (Entrega 1) que entrenó un "
        "clasificador único sobre 993 municipios del país, obteniendo en test "
        "una matriz de confusión equivalente a Precision = 0.43 y Recall = 0.87. "
        "Es un modelo \"alarmista\": detecta la mayoría de los excesos reales "
        "(buen Recall) pero con baja confianza cuando alerta (Precision baja, "
        "~57 % de falsas alarmas). El director del proyecto identificó la causa "
        "principal: mezclar toda la heterogeneidad espacial del país en un solo "
        "modelo, sin variables socioeconómicas que están fuera de alcance.")

    parrafo(doc, "Objetivo cuantitativo del proyecto:", bold=True)
    bullets(doc, [
        "Mejorar Precision en al menos la mitad de los municipios foco "
        "(superar 0.43), manteniendo Recall por encima de 0.60.",
        "Conservar el enfoque metodológico documentado y reproducible.",
        "Entregar artefactos consumibles por un prototipo de dashboard.",
    ])

    parrafo(doc, "Objetivo metodológico:", bold=True)
    bullets(doc, [
        "Entrenar un modelo independiente por municipio foco.",
        "Usar serie temporal continua 2007–2024 (no años representativos sueltos).",
        "Aplicar partición train/test cronológica (último ~20% como test).",
        "Reportar Precision y Recall por municipio, no agregados nacionales.",
    ])

    doc.add_page_break()


def seccion_datos(doc, panel):
    heading(doc, "3. Datos")

    parrafo(doc,
        "El proyecto consume tres fuentes oficiales, integradas en un panel "
        "mensual por municipio:")

    bullets(doc, [
        "SIVIGILA (Instituto Nacional de Salud): casos individuales de dengue "
        "regular (código 210) y grave (código 220), 2007–2024.",
        "Google Earth Engine: variables climáticas mensuales por departamento "
        "—temperatura superficial (MODIS LST), precipitación (CHIRPS), NDVI "
        "(MODIS) y punto de rocío (ERA5-Land) como proxy de humedad.",
        "DANE: proyecciones de población municipal con base CNPV 2018 "
        "(serie unificada 2007–2024 sin escalón entre censos).",
    ])

    n_casos_total = int(panel["casos_total"].sum())
    n_casos_reg = int(panel["casos_regular"].sum())
    n_casos_grave = int(panel["casos_grave"].sum())

    parrafo(doc, "Tabla 1. Tamaño de los datos integrados en el panel.", bold=True)
    agregar_tabla_formateada(doc,
        headers=["Fuente", "Cobertura", "Tamaño"],
        rows=[
            ["SIVIGILA dengue total (en foco)",
             "2007–2024 mensual × 3 municipios",
             f"{n_casos_total:,} casos"],
            ["  · de los cuales regular (210)",
             "—",
             f"{n_casos_reg:,} casos"],
            ["  · de los cuales grave (220)",
             "—",
             f"{n_casos_grave:,} casos"],
            ["GEE clima por departamento",
             "2007–2024 mensual × 33 deptos × 4 variables",
             "7.128 filas climáticas"],
            ["DANE población municipal",
             "2007–2024 anual × 1.123 municipios",
             "Base CNPV 2018 (unificada)"],
            ["Panel mensual procesado (foco)",
             "3 mpios × 18 años × 12 meses",
             "648 filas × 42 columnas"],
        ])

    parrafo(doc,
        "El panel se filtra desde la carga a los 3 municipios foco para reducir "
        "la huella de memoria (de ~6 GB a ~50 MB) y simplificar la iteración. "
        "Los loaders en src/utils.py se mantienen genéricos a nivel país, "
        "permitiendo extensiones futuras a más municipios sin tocar la "
        "infraestructura de datos.")

    doc.add_page_break()


def seccion_metodologia(doc, panel, modelos_xgb):
    heading(doc, "4. Metodología")

    # 4.1 Selección de foco
    heading(doc, "4.1. Selección de los tres municipios foco", level=2)
    parrafo(doc,
        "El director planteó tres criterios válidos para selección espacial: "
        "(a) mayor número absoluto de casos, (b) mayor incidencia por habitante, "
        "(c) interés particular del grupo. Se eligió (b) — mayor incidencia × "
        "100 mil habitantes — porque captura comunidades vulnerables que el "
        "criterio de casos absolutos invisibiliza al estar dominado por ciudades "
        "grandes como Cali o Ibagué.")
    parrafo(doc,
        "Dentro del top 10 nacional por incidencia, se priorizó diversidad "
        "climática y regional para fortalecer la narrativa: tres regiones de "
        "Colombia geográficamente distintas, con dinámicas climáticas y "
        "socioeconómicas contrastantes.")

    rows = []
    for cod, (mun, dpto, region) in DPTO_POR_MPIO.items():
        pob_2024 = int(panel[panel["cod_mpio"] == cod]["poblacion"].dropna().iloc[-1])
        casos = int(panel[panel["cod_mpio"] == cod]["casos_total"].sum())
        rows.append([cod, mun, dpto, region, f"{pob_2024:,}", f"{casos:,}"])
    parrafo(doc, "Tabla 2. Municipios foco con población y casos totales en el panel.", bold=True)
    agregar_tabla_formateada(doc,
        headers=["DIVIPOLA", "Municipio", "Departamento", "Región", "Población 2024", "Casos 2007–2024"],
        rows=rows)

    # 4.2 Población unificada
    heading(doc, "4.2. Población DANE unificada bajo CNPV 2018", level=2)
    parrafo(doc,
        "La versión inicial del archivo de población combinaba proyecciones "
        "base Censo 2005 (2007–2020) con proyecciones base Censo 2018 (2021–2024), "
        "introduciendo un escalón artificial en el límite 2020/2021 de hasta "
        "±25 % por municipio. Se reemplazó por una serie unificada construida "
        "a partir de dos archivos oficiales del DANE, ambos base CNPV 2018: "
        "retroproyecciones municipales 2005–2017 y proyecciones municipales "
        "2018–2042 post COVID-19. La continuidad demográfica se verifica en la "
        "Tabla 3.")

    parrafo(doc, "Tabla 3. Cambios anuales de población 2020→2021 antes y después de unificar la base censal.", bold=True)
    agregar_tabla_formateada(doc,
        headers=["Municipio", "Δ 2020→2021 (antes)", "Δ 2020→2021 (después)"],
        rows=[
            ["Valencia",     "−22 %",  "+0.5 %"],
            ["Fundación",    "+25 %",  "+2.1 %"],
            ["Pueblo Nuevo", "−11 %",  "+1.0 %"],
            ["San Alberto",  "—",      "+2.9 %"],
            ["El Retorno",   "—",      "−0.7 %"],
        ])

    # 4.3 Split temporal
    heading(doc, "4.3. Partición temporal del train y test", level=2)
    parrafo(doc,
        "En series de tiempo no se realiza partición aleatoria — eso produciría "
        "fuga temporal. El test debe ser el periodo más reciente para evaluar "
        "capacidad de pronóstico hacia adelante. Se reservaron los años "
        "2020–2024 (último ~28 % del rango) como test y el resto como train.")

    rows = []
    for cod, (mun, _, _) in DPTO_POR_MPIO.items():
        sub = panel[(panel["cod_mpio"] == cod) & panel["exceso"].notna()]
        tr = sub[sub["ano"] < ANO_TEST_DESDE]
        te = sub[sub["ano"] >= ANO_TEST_DESDE]
        rows.append([
            mun,
            f"{len(tr)} ({tr['exceso'].mean()*100:.0f}% prevalencia)",
            f"{len(te)} ({te['exceso'].mean()*100:.0f}% prevalencia)",
        ])
    parrafo(doc, "Tabla 4. División train/test por municipio (con prevalencia del target).", bold=True)
    agregar_tabla_formateada(doc,
        headers=["Municipio", "Train 2007–2019", "Test 2020–2024"],
        rows=rows)

    # 4.4 Target
    heading(doc, "4.4. Definición del target de exceso epidemiológico", level=2)
    parrafo(doc,
        "El target binario \"exceso\" se define por municipio y por mes "
        "calendario como umbral = máximo entre el percentil 75 histórico (de "
        "casos en ese mismo mes en años anteriores) y un piso de 2 casos. "
        "Un mes cuenta como exceso si los casos observados superan ese umbral. "
        "Esta definición captura la estacionalidad real (lo \"normal\" en julio "
        "no es lo mismo que en enero) y evita falsos positivos en municipios "
        "pequeños donde el percentil histórico puede ser cero por subregistro "
        "o baja transmisión.")

    parrafo(doc,
        "El piso de 2 surgió tras el análisis exploratorio (sección 5.3). Sin "
        "él, en Valencia la prevalencia del target era 57.3 %, contaminada por "
        "alertas espurias en años con pocos casos. Con piso, la prevalencia "
        "baja a 39.6 % y los excesos coinciden con picos visualmente claros.")

    parrafo(doc, "Tabla 5. Comparativo de definiciones evaluadas en el EDA.", bold=True)
    agregar_tabla_formateada(doc,
        headers=["Definición", "Prevalencia Valencia", "Lectura"],
        rows=[
            ["A: percentil 75 sin piso", "57.3 %",
             "Marca como exceso meses con 1–2 casos en 2010–2015. Ruido."],
            ["B: percentil 75 con piso 2  (elegida)", "39.6 %",
             "Los excesos coinciden con los picos visualmente claros."],
            ["C: percentil 90 sin piso", "42.7 %",
             "Todavía contamina los años de bajo reporte."],
        ])

    # 4.5 Features
    heading(doc, "4.5. Features y exclusiones por riesgo de leakage", level=2)
    parrafo(doc,
        "El feature matrix consta de 28 columnas: variables climáticas actuales "
        "(temperatura, precipitación, NDVI, dewpoint), sus rezagos 1, 2 y 3 "
        "meses, sus medias móviles de 3 meses; rezagos 1–3 de casos totales y "
        "de incidencia por 100 mil habitantes; y dos features cíclicas de "
        "estacionalidad (sin/cos del mes).")
    parrafo(doc,
        "Se excluyeron explícitamente variables con riesgo de leakage: "
        "casos_total y sus desagregaciones del mismo mes (definen el target), "
        "hospitalizaciones y fallecidos del mes (información posterior al "
        "evento), umbral_exceso (función directa del histórico), año "
        "(no extrapolable a años nunca vistos por el train) y población "
        "(constante intra-municipio en el train).")

    # 4.6 Modelos
    heading(doc, "4.6. Modelos y estrategia de tuning", level=2)
    parrafo(doc,
        "Se entrenaron tres modelos por municipio:")
    bullets(doc, [
        "Baseline trivial: 1 si los casos del mes anterior superan 2, 0 en otro "
        "caso. Sirve como piso comparativo — \"el modelo memoria pura\".",
        "Regresión Logística regularizada con class_weight='balanced' y "
        "defaults de scikit-learn (regularización L2, C = 1.0).",
        "XGBoost con scale_pos_weight calibrado por municipio y "
        "GridSearchCV(TimeSeriesSplit(5)) sobre un grid pequeño "
        "(n_estimators ∈ {100, 300}, max_depth ∈ {3, 5}, "
        "learning_rate ∈ {0.05, 0.1}); métrica de selección F1.",
    ])
    parrafo(doc,
        "Random Forest fue descartado del comparativo final tras los resultados "
        "de Entrega 1 (Recall ≈ 2.4 % sobre la clase positiva — inservible "
        "para alertar excesos). Deep learning fue descartado a priori por "
        "tamaño de datos (~150 filas de train por municipio).")

    parrafo(doc, "Tabla 6. Hiperparámetros óptimos del XGBoost por municipio.", bold=True)
    rows = []
    for cod, (mun, _, _) in DPTO_POR_MPIO.items():
        bp = modelos_xgb[cod].get("best_params", {})
        rows.append([
            mun,
            str(bp.get("n_estimators", "—")),
            str(bp.get("max_depth", "—")),
            str(bp.get("learning_rate", "—")),
        ])
    agregar_tabla_formateada(doc,
        headers=["Municipio", "n_estimators", "max_depth", "learning_rate"],
        rows=rows)

    doc.add_page_break()


def seccion_eda(doc, panel):
    heading(doc, "5. Análisis exploratorio focalizado")

    # 5.1
    heading(doc, "5.1. Calidad del dato", level=2)
    rows = []
    for cod, (mun, _, _) in DPTO_POR_MPIO.items():
        sub = panel[panel["cod_mpio"] == cod]
        rows.append([
            mun,
            str(len(sub)),
            f"{(sub['casos_total'] == 0).mean() * 100:.0f} %",
            f"{int(sub['casos_total'].sum()):,}",
            f"{sub['casos_total'].max():.0f}",
            f"{int(sub['casos_grave'].sum()):,}",
        ])
    parrafo(doc, "Tabla 7. Resumen de calidad del dato por municipio.", bold=True)
    agregar_tabla_formateada(doc,
        headers=["Municipio", "Meses", "% meses cero", "Casos totales", "Máximo mensual", "Casos graves"],
        rows=rows)
    parrafo(doc,
        "El heatmap de la Figura 1 muestra la distribución anual de casos. "
        "Valencia mantiene niveles muy bajos hasta 2012 (probable subregistro), "
        "despega en 2013 y presenta brotes mayores en 2019 y 2022–2024. "
        "Fundación exhibe el patrón epidemiológico clásico de la región Caribe "
        "con picos anuales nítidos. El Retorno está dominado por dos eventos "
        "extremos (2014 y 2018) y registro escaso en el resto del periodo.")
    parrafo(doc, "Figura 1. Casos anuales por municipio.", bold=True)
    agregar_imagen(doc, GRAPHS_DIR / "01_casos_por_ano_heatmap.png")

    # 5.2
    heading(doc, "5.2. Series temporales por municipio", level=2)
    parrafo(doc,
        "La Figura 2 desagrega las series mensuales. Los marcadores negros "
        "señalan los meses clasificados como exceso bajo la definición elegida "
        "(percentil 75 con piso 2). Se observa cómo el target captura tanto "
        "picos epidémicos grandes como elevaciones estacionales sostenidas.")
    parrafo(doc, "Figura 2. Series temporales mensuales por municipio con excesos marcados.", bold=True)
    agregar_imagen(doc, GRAPHS_DIR / "02_series_temporales.png")

    # 5.3
    heading(doc, "5.3. Comparativo de definiciones de target", level=2)
    parrafo(doc,
        "Se evaluaron tres definiciones del target sobre Valencia (Figura 3). "
        "La variante A (percentil 75 sin piso) genera alertas espurias entre "
        "2010 y 2015, periodo de bajo reporte. La variante C (percentil 90 sin "
        "piso) reduce ese ruido pero aún contamina. La variante B (percentil "
        "75 con piso de 2 casos), elegida, mantiene los picos reales y "
        "elimina los falsos positivos del régimen de bajo reporte.")
    parrafo(doc, "Figura 3. Excesos detectados por las tres definiciones (Valencia).", bold=True)
    agregar_imagen(doc, GRAPHS_DIR / "03_target_comparativo_valencia.png")

    # 5.4
    heading(doc, "5.4. Estacionalidad", level=2)
    parrafo(doc,
        "El boxplot mensual (Figura 4) muestra que Fundación tiene la "
        "estacionalidad más marcada — picos consistentes hacia el último "
        "trimestre del año. Valencia presenta estacionalidad más débil pero "
        "perceptible. El Retorno no muestra patrón estacional claro: la "
        "varianza dentro de cada mes calendario es enorme por los dos eventos "
        "extremos que dominan la distribución.")
    parrafo(doc, "Figura 4. Distribución mensual de casos por municipio.", bold=True)
    agregar_imagen(doc, GRAPHS_DIR / "04_estacionalidad_boxplot.png")

    # 5.5
    heading(doc, "5.5. Correlación clima ↔ casos por lag", level=2)
    parrafo(doc,
        "La Figura 5 muestra correlaciones de Spearman entre casos del mes y "
        "cada variable climática con rezagos 0–3 meses. En Valencia el dewpoint "
        "(humedad) en lag 3 alcanza 0.29 y la temperatura presenta correlación "
        "inversa. En Fundación las correlaciones son débiles (máximo 0.18), "
        "consistente con un patrón epidemiológico bien establecido donde el "
        "predictor dominante es la propia serie. En El Retorno las "
        "correlaciones son tanto débiles como inconsistentes en signo, anticipando "
        "la dificultad del modelado en ese municipio.")
    parrafo(doc, "Figura 5. Correlación Spearman casos ↔ clima por lag.", bold=True)
    agregar_imagen(doc, GRAPHS_DIR / "05_correlacion_clima_lag.png")

    # 5.6
    heading(doc, "5.6. Autocorrelación de casos", level=2)
    parrafo(doc,
        "La Figura 6 muestra el hallazgo más relevante para el modelado: la "
        "transmisión del dengue mensual es fuertemente autoregresiva. En "
        "Valencia y Fundación la autocorrelación lag 1 ronda 0.78 y se "
        "mantiene en 0.45 a lag 12, indicando estacionalidad anual. En El "
        "Retorno la autocorrelación decae rápidamente (0.65 → 0.20), "
        "evidenciando una serie sin memoria estable. Esta estructura "
        "autoregresiva justifica que los rezagos de casos sean features y "
        "anticipa que el modelo va a depender más de su propia serie que del "
        "clima.")
    parrafo(doc, "Figura 6. Autocorrelación de casos por lag.", bold=True)
    agregar_imagen(doc, GRAPHS_DIR / "06_autocorrelacion_casos.png")

    doc.add_page_break()


def seccion_resultados(doc, pred, modelos_xgb):
    heading(doc, "6. Resultados del modelado")

    # 6.1 Tabla maestra
    heading(doc, "6.1. Métricas por municipio y modelo", level=2)
    parrafo(doc,
        "La Tabla 8 consolida Precision, Recall, F1 y Accuracy sobre el test "
        "(2020–2024) para los tres modelos en cada municipio. El comparador "
        "externo es el modelo nacional de Entrega 1 (Precision = 0.43, "
        "Recall = 0.87).")

    rows = []
    for cod, (mun, _, _) in DPTO_POR_MPIO.items():
        for col_pred, modelo in [
            ("pred_baseline", "Baseline (lag1>2)"),
            ("pred_logistic", "Logística"),
            ("pred_xgboost",  "XGBoost"),
        ]:
            m = metricas_modelo(pred, int(cod), col_pred)
            rows.append([
                mun, modelo,
                f"{m['precision']:.2f}",
                f"{m['recall']:.2f}",
                f"{m['f1']:.2f}",
                f"{m['accuracy']:.2f}",
            ])
    parrafo(doc, "Tabla 8. Métricas sobre el test 2020–2024.", bold=True)
    agregar_tabla_formateada(doc,
        headers=["Municipio", "Modelo", "Precision", "Recall", "F1", "Accuracy"],
        rows=rows)

    # 6.2 PR
    heading(doc, "6.2. Comparativo visual Precision vs Recall", level=2)
    parrafo(doc,
        "La Figura 7 ubica a cada modelo en el plano Precision–Recall. Las "
        "líneas grises punteadas marcan los valores del modelo nacional "
        "(Precision = 0.43, Recall = 0.87). Los modelos foco para Valencia y "
        "Fundación se ubican claramente por encima en Precision con Recall "
        "comparable.")
    parrafo(doc, "Figura 7. Precision vs Recall sobre el test 2020–2024.", bold=True)
    agregar_imagen(doc, GRAPHS_DIR / "07_precision_recall.png")

    # 6.3 Matrices
    heading(doc, "6.3. Matrices de confusión", level=2)
    parrafo(doc,
        "La Figura 8 muestra las matrices de confusión por municipio y modelo. "
        "Para vigilancia epidemiológica, los falsos negativos (FN, brotes no "
        "alertados) son operativamente más costosos que los falsos positivos "
        "(FP, alertas sin brote): un FN implica preparación sanitaria "
        "deficiente, un FP implica una campaña preventiva que pudo haberse "
        "evitado.")
    parrafo(doc, "Figura 8. Matrices de confusión.", bold=True)
    agregar_imagen(doc, GRAPHS_DIR / "07_matrices_confusion.png")

    # 6.4 Feature importance
    heading(doc, "6.4. Importancia global de features (XGBoost)", level=2)
    parrafo(doc,
        "La Figura 9 muestra la importancia global de las 10 features más "
        "relevantes para XGBoost en cada municipio. El patrón es consistente: "
        "los rezagos de casos o incidencia dominan, y las features climáticas "
        "aparecen como moduladores secundarios.")
    parrafo(doc, "Figura 9. Top 10 features XGBoost por municipio.", bold=True)
    agregar_imagen(doc, GRAPHS_DIR / "07_feature_importance.png")

    parrafo(doc, "Tabla 9. Feature más importante en cada municipio (XGBoost).", bold=True)
    rows = []
    for cod, (mun, _, _) in DPTO_POR_MPIO.items():
        m = modelos_xgb[cod]
        importancias = pd.Series(m["model"].feature_importances_,
                                 index=m["features"]).sort_values(ascending=False)
        top1 = importancias.index[0]
        rows.append([mun, top1, f"{importancias.iloc[0]:.3f}"])
    agregar_tabla_formateada(doc,
        headers=["Municipio", "Feature dominante", "Importancia"],
        rows=rows)

    # 6.5 SHAP
    heading(doc, "6.5. Interpretabilidad local con SHAP", level=2)
    parrafo(doc,
        "Las Figuras 10–12 muestran los valores SHAP del modelo XGBoost para "
        "cada municipio sobre el test. A diferencia de la importancia global, "
        "SHAP descompone cada predicción individual mostrando magnitud y signo "
        "del impacto de cada feature. Color rojo indica valor alto del feature "
        "en esa observación; azul, valor bajo.")

    parrafo(doc, "Figura 10. SHAP — Valencia.", bold=True)
    agregar_imagen(doc, GRAPHS_DIR / "07b_shap_23855.png")
    parrafo(doc,
        "En Valencia los rezagos de casos e incidencia dominan: valores altos "
        "del mes pasado empujan fuertemente hacia exceso (SHAP positivo). El "
        "dewpoint y la temperatura aparecen como segunda capa de señal.")

    parrafo(doc, "Figura 11. SHAP — Fundación.", bold=True)
    agregar_imagen(doc, GRAPHS_DIR / "07b_shap_47288.png")
    parrafo(doc,
        "Fundación muestra el efecto autoregresivo más intenso: casos_total_lag1 "
        "produce SHAP de hasta +4 (escala log-odds). La precipitación en lag 3 "
        "alta tiende a SHAP negativo, posible captura de un ciclo \"exceso de "
        "lluvia → renovación de criaderos → caída temporal\".")

    parrafo(doc, "Figura 12. SHAP — El Retorno.", bold=True)
    agregar_imagen(doc, GRAPHS_DIR / "07b_shap_95025.png")
    parrafo(doc,
        "En El Retorno los patrones SHAP son más débiles e inconsistentes "
        "biológicamente: NDVI alto se asocia con SHAP negativo, contrario a la "
        "intuición ecológica. Esto refuerza la limitación documentada — el "
        "modelo captura señal pero no es robusta.")

    doc.add_page_break()


def seccion_conclusiones(doc, pred):
    heading(doc, "7. Conclusiones y limitaciones")

    # 7.1 Cumplimiento
    heading(doc, "7.1. Cumplimiento del objetivo cuantitativo", level=2)
    parrafo(doc,
        "El objetivo cuantitativo (Precision > 0.43 manteniendo Recall > 0.60) "
        "se evalúa en la Tabla 10 con el mejor modelo entrenado por municipio.")

    rows = []
    cumple = 0
    for cod, (mun, _, _) in DPTO_POR_MPIO.items():
        m = metricas_modelo(pred, int(cod), "pred_xgboost")
        ok = m["precision"] > 0.43 and m["recall"] > 0.60
        cumple += int(ok)
        rows.append([
            mun,
            f"{m['precision']:.2f}",
            f"{m['recall']:.2f}",
            "Sí" if ok else "No",
        ])
    agregar_tabla_formateada(doc,
        headers=["Municipio", "Precision (XGBoost)", "Recall (XGBoost)", "Cumple"],
        rows=rows)

    parrafo(doc,
        f"El objetivo se cumple en {cumple} de 3 municipios. En el otro caso "
        f"(El Retorno) los modelos entrenados no superan al baseline trivial; "
        f"el baseline sí cumple con Precision = 0.55 y Recall = 0.79, lo que "
        f"sugiere que para municipios con series escasas y dominadas por "
        f"outliers, la información de \"qué pasó el mes anterior\" es "
        f"prácticamente toda la señal disponible y los modelos más complejos "
        f"sobreajustan ruido.")

    # 7.2 Hallazgo: baseline competitivo
    heading(doc, "7.2. Hallazgo metodológico: la autoregresión domina", level=2)
    parrafo(doc,
        "El baseline trivial (\"casos del mes pasado > 2\") es notablemente "
        "competitivo en los tres municipios: F1 baseline 0.88 en Valencia, "
        "0.78 en Fundación, 0.65 en El Retorno. XGBoost añade entre 1 y 14 "
        "puntos porcentuales en Precision sin sacrificar Recall — un valor "
        "incremental real pero menor que el esperado intuitivamente. La "
        "razón estructural es que la dinámica de transmisión del dengue "
        "mensual está dominada por su propia inercia temporal; el clima y "
        "otras covariables modulan secundariamente. Este hallazgo es honesto "
        "y útil: un usuario que solo tenga acceso a SIVIGILA del mes pasado "
        "ya tiene la mayor parte de la información predictiva.")

    # 7.3 Limitaciones
    heading(doc, "7.3. Limitaciones reconocidas", level=2)
    bullets(doc, [
        "n = 3 municipios. La metodología se ejemplifica pero no se valida "
        "estadísticamente sobre un universo amplio.",
        "El Retorno: ningún modelo entrenado supera al baseline. La aproximación "
        "regional tiene un piso de aplicabilidad en series con eventos "
        "extremos dominantes.",
        "Aporte climático moderado. Las correlaciones lineales casos ↔ clima "
        "no superan 0.29; el clima aparece como modulador secundario, no "
        "como driver principal. XGBoost extrae interacciones no lineales, "
        "pero el grueso de la señal sigue siendo la autoregresión.",
        "El umbral de probabilidad de clasificación se dejó en 0.5 (defecto). "
        "Una calibración por F1 óptimo sobre validación podría mover Precision "
        "y Recall pero queda como ajuste del consumidor final (dashboard).",
        "Las variables climáticas del mes corriente forman parte del feature "
        "set. Como el clima es exógeno al target, no constituye leakage "
        "estadístico; pero en despliegue en tiempo real, su disponibilidad "
        "depende de la frescura del dato climático.",
    ])

    doc.add_page_break()


def seccion_sensibilidad(doc, sens):
    """Análisis de sensibilidad del umbral del target (P75 piso 0-3 vs P90)."""
    heading(doc, "7.4. Análisis de sensibilidad del umbral del target", level=2)

    parrafo(doc,
        "La definición del target (percentil 75 histórico con piso de 2 casos, "
        "decisión D12) se eligió por inspección visual del EDA. Para verificar "
        "que los resultados no dependen fuertemente de esa elección puntual, "
        "se re-evaluó el modelo XGBoost manteniendo los hiperparámetros "
        "óptimos ya encontrados y variando la definición del target en una "
        "grilla: percentil ∈ {75, 90} × piso ∈ {0, 1, 2, 3}. Ocho "
        "combinaciones × 3 municipios = 24 re-entrenamientos.")

    # Tabla resumen con F1 para todas las combinaciones P75
    parrafo(doc, "Tabla 12. F1 sobre el test 2020–2024 al variar el piso (XGBoost con percentil 75).", bold=True)
    p75 = sens[(sens["percentil"] == 75) & (~sens["degenerado"])]
    rows = []
    for cod, (mun, _, _) in DPTO_POR_MPIO.items():
        sub = p75[p75["municipio"] == mun].sort_values("piso")
        row = [mun]
        for piso in [0, 1, 2, 3]:
            r = sub[sub["piso"] == piso]
            if r.empty:
                row.append("—")
            else:
                marker = " ★" if piso == 2 else ""
                row.append(f"{r.iloc[0]['f1']:.2f}{marker}")
        rows.append(row)
    agregar_tabla_formateada(doc,
        headers=["Municipio", "Piso 0", "Piso 1", "Piso 2 (★ usado)", "Piso 3"],
        rows=rows)

    # Comparativo P75 vs P90 con piso 2
    parrafo(doc, "Tabla 13. P75 contra P90 (manteniendo piso 2): impacto en Precision/Recall.", bold=True)
    rows = []
    for cod, (mun, _, _) in DPTO_POR_MPIO.items():
        for perc in [75, 90]:
            r = sens[(sens["percentil"] == perc) & (sens["piso"] == 2)
                     & (sens["municipio"] == mun) & (~sens["degenerado"])]
            if not r.empty:
                rows.append([
                    mun, f"P{perc}",
                    f"{r.iloc[0]['precision']:.2f}",
                    f"{r.iloc[0]['recall']:.2f}",
                    f"{r.iloc[0]['f1']:.2f}",
                ])
    agregar_tabla_formateada(doc,
        headers=["Municipio", "Percentil", "Precision", "Recall", "F1"],
        rows=rows)

    parrafo(doc, "Figura 13. Sensibilidad de XGBoost al umbral del target (Precision/Recall/F1 vs piso).", bold=True)
    agregar_imagen(doc, GRAPHS_DIR / "08_sensibilidad_target.png")

    # Compute variation ranges for the conclusion
    variaciones = {}
    for cod, (mun, _, _) in DPTO_POR_MPIO.items():
        sub = p75[p75["municipio"] == mun]
        if not sub.empty:
            variaciones[mun] = sub["f1"].max() - sub["f1"].min()

    parrafo(doc,
        f"Lectura: con percentil 75, el F1 varía entre piso 0 y piso 3 en "
        f"{variaciones.get('Valencia', 0):.2f} puntos en Valencia, "
        f"{variaciones.get('Fundación', 0):.2f} en Fundación y "
        f"{variaciones.get('El Retorno', 0):.2f} en El Retorno. La elección "
        f"de piso 2 se ubica dentro del rango razonable: en Valencia el F1 "
        f"está en plateau (piso 2 y 3 dan idéntico 0.88); en Fundación piso 3 "
        f"daría una leve mejora adicional pero piso 2 ya cumple el objetivo del "
        f"proyecto; en El Retorno piso 2 está cerca del óptimo dentro de la "
        f"grilla (0.53). El percentil 90 es claramente inferior en los 3 "
        f"municipios — su F1 cae 0.20–0.30 puntos respecto a P75 — porque "
        f"genera modelos demasiado alarmistas (Recall alto pero Precision "
        f"deprimida, similar al modelo nacional original). Conclusión: los "
        f"resultados del proyecto son estables a la elección del umbral dentro "
        f"del régimen razonable; la decisión de usar P75 con piso 2 está "
        f"respaldada por evidencia cuantitativa y no es producto de un ajuste "
        f"selectivo sobre la métrica final.")

    doc.add_page_break()


def seccion_artefactos(doc):
    heading(doc, "8. Artefactos y reproducibilidad")

    parrafo(doc,
        "Todo el código y los datos derivados están versionados en el repositorio "
        "del proyecto. La Tabla 11 resume los artefactos clave.")

    parrafo(doc, "Tabla 11. Artefactos del proyecto.", bold=True)
    agregar_tabla_formateada(doc,
        headers=["Tipo", "Ruta", "Descripción"],
        rows=[
            ["Notebook EDA", "notebooks/06_eda_foco.ipynb", "EDA focalizado con figuras embebidas"],
            ["Notebook modelado", "notebooks/07_modelado.ipynb", "Entrenamiento + métricas + SHAP"],
            ["Panel mensual", "data/processed/panel_municipal_mensual.parquet", "648 filas × 42 columnas"],
            ["Predicciones test", "data/processed/predicciones_test.csv", "Salidas de los 3 modelos × 3 mpios"],
            ["Modelos XGBoost", "models/{23855,47288,95025}_xgboost.joblib", "Uno por municipio"],
            ["Modelos Logística", "models/{23855,47288,95025}_logistic.joblib", "Uno por municipio"],
            ["Decisiones", "docs/decisiones_proyecto.md", "17 decisiones D1–D17 con justificación"],
            ["Reunión origen", "docs/reuniones/2026-04-24_seguimiento_carlos.md", "Acta director del proyecto"],
            ["Gráficos", "results_graphs/foco/*.png", "12 figuras del EDA y modelado"],
            ["Bundle dashboard", "scripts/exportar_artefactos_dashboard.py", "Empaqueta artefactos para el repo del dashboard"],
        ])

    parrafo(doc,
        "El pipeline es reproducible end-to-end ejecutando los scripts en "
        "scripts/ en el orden documentado por las decisiones D1–D17. Los "
        "modelos serializados están listos para consumo desde el dashboard de "
        "alerta temprana (repo dashboard_dengue).")


# ============================================================================
# Main
# ============================================================================
def main():
    print("Cargando datos…")
    panel = pd.read_parquet(PANEL_PATH)
    pred = pd.read_csv(PRED_PATH)
    pred["cod_mpio"] = pred["cod_mpio"].astype(int)  # consistente con MUNICIPIOS_FOCO usado como int
    panel["cod_mpio"] = panel["cod_mpio"].astype(str)  # como string para el filtrado por foco
    sens = pd.read_csv(SENS_PATH) if SENS_PATH.exists() else None
    modelos_xgb = {
        cod: joblib.load(MODELS_DIR / f"{cod}_xgboost.joblib")
        for cod in MUNICIPIOS_FOCO
    }
    print(f"  panel:        {panel.shape[0]} filas × {panel.shape[1]} cols")
    print(f"  predicciones: {pred.shape[0]} filas × {pred.shape[1]} cols")
    print(f"  sensibilidad: {sens.shape[0] if sens is not None else '—'} filas")
    print(f"  modelos XGB:  {len(modelos_xgb)}")

    print("Construyendo documento…")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    agregar_portada(doc, TEXTOS)
    seccion_resumen(doc, pred)
    seccion_contexto(doc)
    seccion_datos(doc, panel)
    seccion_metodologia(doc, panel, modelos_xgb)
    seccion_eda(doc, panel)
    seccion_resultados(doc, pred, modelos_xgb)
    seccion_conclusiones(doc, pred)
    if sens is not None:
        seccion_sensibilidad(doc, sens)
    seccion_artefactos(doc)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    fecha = datetime.now().strftime("%Y-%m-%d")
    out_path = OUT_DIR / f"reporte_proyecto_dengue_{fecha}.docx"
    doc.save(out_path)
    print(f"\n✅ Reporte generado: {out_path}")
    print(f"   Tamaño: {out_path.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
